import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def insert_hr(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr,
        'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
        'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
        'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
        'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
        'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
        'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
        'w:pPrChange'
    )
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F497D')
    pBdr.append(bottom)

def create_resume():
    folder_name = "Curriculo"
    if not os.path.exists(folder_name):
        os.makedirs(folder_name)

    doc = Document()
    
    # Page setup
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Base style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)
    
    # Colors
    DARK_BLUE = RGBColor(31, 73, 125)
    GRAY = RGBColor(89, 89, 89)

    # Header
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_p.add_run("RICARDO MACEDO MARTINES")
    name_run.bold = True
    name_run.font.size = Pt(18)
    name_run.font.color.rgb = DARK_BLUE

    role_p = doc.add_paragraph()
    role_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    role_run = role_p.add_run("Planejamento Estratégico & Transformação Digital | Processos, Governança & Inteligência de Dados")
    role_run.italic = True
    role_run.font.size = Pt(12)
    role_run.font.color.rgb = GRAY
    
    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_run = contact_p.add_run("(11) 95421-0088 | ricardo.martines@live.com | linkedin.com/in/eng-ricardomartines/ | São Paulo - SP")
    contact_run.font.size = Pt(10)
    contact_p.paragraph_format.space_after = Pt(20)

    def add_section_header(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = DARK_BLUE
        insert_hr(p)
    
    # RESUMO PROFISSIONAL
    add_section_header("RESUMO PROFISSIONAL")
    
    p = doc.add_paragraph("Sou Engenheiro de Produção com MBA em Estratégia de Negócios e construí uma trajetória de mais de 20 anos liderando frentes de eficiência operacional, redesenho de processos e governança de projetos em grandes players de mercado. Hoje, atuo focado na Coordenação de Transformação Digital, posicionando-me como um ")
    p.add_run("Business Translator").bold = True
    p.add_run(": o elo de confiança que traduz as dores estratégicas da Alta Liderança (C-Level) em soluções de tecnologia eficientes, lógicas e escaláveis para os times de desenvolvimento (TI).")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p = doc.add_paragraph("Minha especialidade é estruturar e gerenciar esteiras operacionais de alta performance, eliminando o desperdício oculto e garantindo o cumprimento rigoroso de prazos e metas (SLA). Tenho sólida experiência em liderar processos de mudança organizacional, atuando no Planejamento Estratégico em parceria com o C-Level para garantir a escalabilidade do negócio.")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p = doc.add_paragraph("Não olho para a tecnologia apenas como inovação conceitual, mas como um investimento focado em margem financeira e eficiência mensurável. Meu histórico comprova essa entrega de valor: liderei frentes que geraram ")
    p.add_run("R$ 3 milhões em savings operacionais em apenas 7 meses").bold = True
    p.add_run(" na minha atual cadeira, além de ter capturado ")
    p.add_run("R$ 258 milhões em sinergias").bold = True
    p.add_run(" de processos em cenários complexos de fusões e aquisições (M&A).")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # PRINCIPAIS COMPETÊNCIAS
    add_section_header("PRINCIPAIS COMPETÊNCIAS")
    
    def add_bullet(bold_text, normal_text):
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(bold_text).bold = True
        p.add_run(normal_text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        return p

    add_bullet("Motores de Hiperautomação: ", "Plataforma n8n (iPaaS) e ecossistema Microsoft Power Platform (Power Automate, Power Apps).")
    add_bullet("Inteligência de Dados e Desenvolvimento: ", "Power BI Avançado, Lógica de Programação, JavaScript, Java e modelagem de bancos de dados SQL.")
    add_bullet("Governança e Processos: ", "Business Process Redesign (BPR), certificação Lean Six Sigma Green Belt, Ciclo PDCA, Frameworks Ágeis (Scrum / Kanban), Gestão de Mudança (Change Management), Governança Corporativa e Gestão de SLAs Corporativos.")

    # EXPERIÊNCIA PROFISSIONAL
    add_section_header("EXPERIÊNCIA PROFISSIONAL")

    def add_job_header(company, dates, role):
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        table.columns[0].width = Inches(5.0)
        table.columns[1].width = Inches(1.8)
        row = table.rows[0]
        p_left = row.cells[0].paragraphs[0]
        run_c = p_left.add_run(company)
        run_c.bold = True
        run_c.font.color.rgb = DARK_BLUE
        
        p_right = row.cells[1].paragraphs[0]
        p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run_d = p_right.add_run(dates)
        run_d.bold = True
        run_d.font.color.rgb = GRAY
        
        p_role = doc.add_paragraph()
        p_role.paragraph_format.space_before = Pt(2)
        p_role.paragraph_format.space_after = Pt(6)
        run_r = p_role.add_run(role)
        run_r.italic = True

    add_job_header("Metal Alfa Ltda. (Telecomunicações, Metalurgia e Infraestrutura)", "2025 - Momento", "Coordenador de Planejamento Estratégico e Transformação Digital")
    add_bullet("Mapeamento & BPR: ", "Liderança na reestruturação e governança de macroprocessos ponta a ponta (produção, obras e logística), implementando metodologias de gestão que eliminaram gargalos operacionais e garantiram previsibilidade na esteira de entregas da companhia.")
    add_bullet("Arquitetura de Hiperautomação: ", "Direcionamento estratégico para a implantação de motores de Hiperautomação e Inteligência de Dados, convertendo processos manuais analógicos em uma operação digital escalável e auditável.")
    add_bullet("Governança Executiva: ", "Estruturação de uma cultura puramente baseada em dados, criando rituais de controle com dashboards no Power BI para suporte a decisões críticas da Diretoria.")
    add_bullet("Impacto Financeiro: ", "Geração de R$ 3 milhões em savings operacionais em apenas 7 meses de atuação através das soluções e automações implementadas.")

    add_job_header("Winity S.A.", "2024", "PMO de Processos e Projetos Estratégicos (Projeto Vale)")
    add_bullet("Business Translator: ", "Atuação como elo estratégico entre frentes multidisciplinares (ESG, Comercial, Qualidade) e a Diretoria Executiva, unificando os requisitos de negócio em metas claras.")
    add_bullet("Mitigação de Riscos: ", "Mapeamento de macroprocessos críticos, desenho de fluxos preventivos e implementação de frameworks ágeis (Scrum/Kanban) para dar previsibilidade às entregas e contratos.")

    add_job_header("Claro S.A.", "2020 - 2023", "Analista de Projetos e Processos Sênior")
    add_bullet("Sinergia em M&A: ", "Liderança analítica em estudos de fusões e aquisições (integração Nextel/Claro e Claro/Oi), desenhando os novos fluxos operacionais unificados.")
    add_bullet("Automação de Indicadores: ", "Modelagem e automatização de matrizes de KPIs para medição automatizada de marcos críticos (Milestones) de projetos de infraestrutura de rede.")
    add_bullet("Impacto Financeiro: ", "Atuação direta na captura de um saving acumulado de R$ 258 milhões por meio da otimização e eliminação de gaps de processos pós-fusão.")

    add_job_header("Nextel S.A.", "2018 - 2019", "Engenheiro de Implantação e Eficiência Financeira")
    add_bullet("Automação de Relatórios: ", "Inovação nos rituais de reporte operacional (Rollout) via desenvolvimento de macros e rotinas digitais, substituindo controles manuais por visibilidade de KPIs em tempo real para o C-Level.")
    add_bullet("Impacto Operacional: ", "Redução drástica de 16 horas semanais (HH) de trabalho operacional repetitivo da equipe e garantia de 100% de aderência ao budget de CAPEX/OPEX.")

    add_job_header("Atento S.A. & Quadra Contact Center", "2015 - 2018", "Coordenador / Especialista em Planejamento Operacional")
    add_bullet("Otimização Estrutural: ", "Aplicação de ciclos PDCA e modelagem de processos de atendimento e backoffice, gerando uma redução de 7% nos custos diretos com HH e elevando os indicadores de satisfação contratual (SLA).")

    p_nota = doc.add_paragraph()
    p_nota.paragraph_format.space_before = Pt(6)
    run_nota = p_nota.add_run("(Nota: Histórico inicial incluindo atuação em manufatura e PCP na SKF Brasil e demais empresas consolidado no perfil do LinkedIn).")
    run_nota.italic = True
    run_nota.font.size = Pt(9)
    run_nota.font.color.rgb = GRAY

    # FORMAÇÃO ACADÊMICA E CERTIFICAÇÕES
    add_section_header("FORMAÇÃO ACADÊMICA E CERTIFICAÇÕES")
    
    add_bullet("MBA em Estratégia de Negócios", " - Faculdades Metropolitanas Unidas (FMU)")
    add_bullet("Graduação em Engenharia de Produção", " - Universidade Anhembi Morumbi")
    add_bullet("Formação Tech & Automação: ", "n8n Avançado / iPaaS (2025) | Microsoft Power Platform (2024) | Lógica de Programação, JavaScript & Java (2025)")
    add_bullet("Metodologias de Processos: ", "Certificação Lean Six Sigma Green Belt | Frameworks Ágeis (Scrum, Kanban)")

    file_path = os.path.join(folder_name, "Curriculo_Ricardo_Martines.docx")
    doc.save(file_path)
    print(f"Document saved to {file_path}")

if __name__ == '__main__':
    create_resume()
